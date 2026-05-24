"""
Word (DOCX) to PDF converter.
Uses python-docx to read the document and reportlab to render the PDF.
Extracts paragraphs, tables, and basic formatting.
"""
import os
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
)
from reportlab.lib import colors


def convert(input_path: str, output_path: str) -> str:
    """
    Convert a DOCX file to PDF format.
    
    Args:
        input_path: Absolute path to the source DOCX file.
        output_path: Absolute path for the output PDF file.
    
    Returns:
        The output_path on success.
    
    Raises:
        ValueError: If input file doesn't exist or is not a DOCX.
        RuntimeError: If conversion fails.
    """
    if not os.path.exists(input_path):
        raise ValueError(f"Input file not found: {input_path}")
    
    if not input_path.lower().endswith(('.docx', '.doc')):
        raise ValueError("Input file must be a DOCX.")
    
    try:
        doc = Document(input_path)
        
        # Build PDF
        pdf = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=25 * mm,
            rightMargin=25 * mm,
            topMargin=25 * mm,
            bottomMargin=25 * mm
        )
        
        styles = getSampleStyleSheet()
        
        # Custom styles for different heading levels
        custom_styles = {
            'DocHeading1': ParagraphStyle(
                'DocHeading1', parent=styles['Heading1'],
                fontSize=18, spaceAfter=12, textColor=colors.black
            ),
            'DocHeading2': ParagraphStyle(
                'DocHeading2', parent=styles['Heading2'],
                fontSize=15, spaceAfter=10, textColor=colors.black
            ),
            'DocHeading3': ParagraphStyle(
                'DocHeading3', parent=styles['Heading3'],
                fontSize=13, spaceAfter=8, textColor=colors.black
            ),
            'DocBody': ParagraphStyle(
                'DocBody', parent=styles['Normal'],
                fontSize=11, leading=15, spaceAfter=6, textColor=colors.black
            ),
        }
        
        story = []
        
        for element in doc.element.body:
            tag = element.tag.split('}')[-1]  # Remove namespace
            
            if tag == 'p':
                # It's a paragraph
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                
                if para is None:
                    continue
                
                text = para.text.strip()
                if not text:
                    story.append(Spacer(1, 6))
                    continue
                
                # Determine style from paragraph style name
                style_name = (para.style.name or '').lower()
                
                if 'heading 1' in style_name or 'title' in style_name:
                    style = custom_styles['DocHeading1']
                elif 'heading 2' in style_name:
                    style = custom_styles['DocHeading2']
                elif 'heading 3' in style_name:
                    style = custom_styles['DocHeading3']
                else:
                    style = custom_styles['DocBody']
                
                # Handle bold/italic at run level
                formatted_text = ""
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue
                    # Escape XML special chars for reportlab
                    run_text = (run_text
                        .replace('&', '&amp;')
                        .replace('<', '&lt;')
                        .replace('>', '&gt;'))
                    
                    if run.bold and run.italic:
                        formatted_text += f"<b><i>{run_text}</i></b>"
                    elif run.bold:
                        formatted_text += f"<b>{run_text}</b>"
                    elif run.italic:
                        formatted_text += f"<i>{run_text}</i>"
                    else:
                        formatted_text += run_text
                
                if not formatted_text:
                    formatted_text = (text
                        .replace('&', '&amp;')
                        .replace('<', '&lt;')
                        .replace('>', '&gt;'))
                
                story.append(Paragraph(formatted_text, style))
            
            elif tag == 'tbl':
                # It's a table
                for tbl in doc.tables:
                    if tbl._element is element:
                        table_data = []
                        for row in tbl.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                cell_text = (cell_text
                                    .replace('&', '&amp;')
                                    .replace('<', '&lt;')
                                    .replace('>', '&gt;'))
                                row_data.append(
                                    Paragraph(cell_text, custom_styles['DocBody'])
                                )
                            table_data.append(row_data)
                        
                        if table_data:
                            t = Table(table_data)
                            t.setStyle(TableStyle([
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.9)),
                                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                                ('PADDING', (0, 0), (-1, -1), 6),
                            ]))
                            story.append(t)
                            story.append(Spacer(1, 12))
                        break
        
        if not story:
            # If no content extracted, add a message
            story.append(Paragraph("No content could be extracted.", custom_styles['DocBody']))
        
        # Build the PDF — this writes fully to disk before returning
        pdf.build(story)
        
    except Exception as e:
        raise RuntimeError(f"DOCX to PDF conversion failed: {str(e)}")
    
    # Verify output
    if not os.path.exists(output_path):
        raise RuntimeError("Conversion produced no output file.")
    if os.path.getsize(output_path) < 100:
        raise RuntimeError("Conversion produced an empty or corrupt file.")
    
    return output_path
